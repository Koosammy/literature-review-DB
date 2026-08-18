"""DOCX table and figure extraction.

Strategy: DOCX has no fixed page layout, so there is no way to "crop the
real page" the way the PDF engine does. Instead we convert the document to
PDF with a headless LibreOffice (which renders it exactly as Word would)
and reuse :class:`PdfExtractor` on the result -- this gives DOCX tables the
same pixel-accurate crops as PDF tables, including merged cells and
multi-line headers that a from-scratch re-typeset would get wrong.

Embedded images are read directly from the .docx zip (``word/media/*``)
rather than from the converted PDF, since that avoids an extra lossy
re-encode. If LibreOffice isn't installed, we fall back to a pure
python-docx table reader (cell text only, no pixel-perfect crop) so the
tool still works, just with lower table fidelity.
"""

from __future__ import annotations

import logging
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import List, Optional

import pandas as pd

from .models import ExtractedFigure, ExtractedTable, ExtractionResult
from .pdf_engine import PdfExtractor
from .utils import sha256_bytes

logger = logging.getLogger(__name__)


def _docx_tc_text(tc) -> str:
    """A table cell's own text (its paragraphs joined with "\\n"), scoped to
    just this ``<w:tc>`` -- unlike python-docx's ``row.cells[i].text``, this
    does NOT inherit a vertically-merged parent cell's text, which is what
    lets the header-merge detection below tell "this cell has nothing of
    its own to add" apart from "this cell is empty"."""
    from docx.oxml.ns import qn

    lines = []
    for p in tc.findall(qn("w:p")):
        lines.append("".join(t.text or "" for t in p.iter(qn("w:t"))))
    return "\n".join(lines).strip()


_NO_FILL = {None, "auto", "FFFFFF"}


def _is_light_hex(hex_color: str) -> bool:
    try:
        r, g, b = (int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
    except (ValueError, IndexError):
        return True
    # Standard relative-luminance approximation for text contrast.
    return (0.299 * r + 0.587 * g + 0.114 * b) > 150


def _docx_table_grid(table) -> List[List[Optional[tuple]]]:
    """Expand a table into a full row x column grid of (own_text, span,
    vmerge, fill) tuples, duplicating a horizontally-merged (gridSpan)
    cell's own text into each column it spans. ``vmerge`` is ``"continue"``
    for a cell that's the tail of a vertical merge (Word leaves these with
    no text of their own -- the real content lives on the "restart" cell
    above). ``fill`` is the cell's background color (a "RRGGBB" hex string)
    when Word records one, else ``None`` -- some tables (e.g. a Gantt-style
    schedule) mark their data with cell shading and no text at all, so this
    is the only way to recover that content."""
    from docx.oxml.ns import qn

    n_cols = len(table.columns)
    grid_rows = []
    for row in table.rows:
        grid = [None] * n_cols
        col = 0
        for tc in row._tr.findall(qn("w:tc")):
            text = _docx_tc_text(tc)
            span = 1
            vmerge = None
            fill = None
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    span = int(gs.get(qn("w:val")) or 1)
                vm = tcPr.find(qn("w:vMerge"))
                if vm is not None:
                    vmerge = vm.get(qn("w:val")) or "continue"
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    raw_fill = shd.get(qn("w:fill"))
                    if raw_fill and raw_fill.upper() not in _NO_FILL:
                        fill = raw_fill.upper()
            for k in range(span):
                if col + k < n_cols:
                    grid[col + k] = (text, span, vmerge, fill)
            col += span
        grid_rows.append(grid)
    return grid_rows


def _docx_table_header_and_body(table) -> tuple:
    """Read a python-docx Table into (header, body_rows), correctly merging
    a header that spans more than one physical table row.

    A merged parent header cell (e.g. one "Delivery" cell horizontally
    spanning two columns via Word's gridSpan) is ambiguous on its own: both
    spanned columns show the same label. Real documents resolve that
    ambiguity with a second header row giving each spanned column its own
    sub-label (e.g. "Skilled" / "Unskilled") -- and leave every already-
    unambiguous column blank on that row, whether or not Word recorded an
    explicit vMerge for it (it often doesn't; a plain blank cell is just as
    common as an explicit vertical merge). So the rule used here is
    structural, not markup-dependent: a row extends the header for as long
    as every already-unambiguous column stays blank on it while at least
    one still-ambiguous (gridSpan > 1) column gets its own text -- the
    moment a resolved column has its own content, that row is real data.
    """
    grid_rows = _docx_table_grid(table)
    if not grid_rows:
        return [], [], []

    n_cols = len(grid_rows[0])
    header_parts = [[grid_rows[0][c][0] if grid_rows[0][c] else ""] for c in range(n_cols)]
    resolved = [not grid_rows[0][c] or grid_rows[0][c][1] == 1 for c in range(n_cols)]

    header_end = 1
    r = 1
    while r < len(grid_rows) and not all(resolved):
        row = grid_rows[r]
        ok = True
        any_new = False
        for c in range(n_cols):
            text = row[c][0] if row[c] else ""
            if resolved[c]:
                if text:
                    ok = False
                    break
            elif text:
                any_new = True
        if not ok or not any_new:
            break
        for c in range(n_cols):
            if not resolved[c]:
                text = row[c][0] if row[c] else ""
                if text:
                    header_parts[c].append(text)
                    resolved[c] = True
        header_end = r + 1
        r += 1

    header = ["\n".join(part for part in parts if part) for parts in header_parts]

    # A vertically-merged data cell (e.g. one category label spanning
    # several rows) has empty text of its own on every row but the first;
    # inherit from directly above only for those, not for cells that are
    # just blank on their own merits (most blanks in these tables mean
    # "not applicable", not "same as above").
    body_rows = []
    body_fills: List[List[Optional[str]]] = []
    last_seen = [""] * n_cols
    for r in range(header_end, len(grid_rows)):
        row_vals = []
        row_fills = []
        for c in range(n_cols):
            cell = grid_rows[r][c]
            text = cell[0] if cell else ""
            vmerge = cell[2] if cell else None
            fill = cell[3] if cell else None
            if not text and vmerge == "continue":
                text = last_seen[c]
            last_seen[c] = text
            # A cell marked only by background shading (no text at all --
            # e.g. a schedule's "active this month" marker) still needs
            # *some* content, or it reads as missing data and the table as
            # a whole can look empty enough to be mistaken for noise.
            if not text and fill:
                text = "●"
            row_vals.append(text)
            row_fills.append(fill)
        body_rows.append(row_vals)
        body_fills.append(row_fills)

    return header, body_rows, body_fills


class DocxExtractor:
    def __init__(
        self,
        *,
        pdf_extractor: Optional[PdfExtractor] = None,
        min_image_bytes: int = 4000,
        libreoffice_timeout: int = 90,
    ) -> None:
        self.pdf_extractor = pdf_extractor or PdfExtractor()
        self.min_image_bytes = min_image_bytes
        self.libreoffice_timeout = libreoffice_timeout

    def extract(self, data: bytes) -> ExtractionResult:
        result = ExtractionResult()

        media_images = self._extract_media_images(data)
        result.figures.extend(media_images)

        pdf_bytes = self._convert_to_pdf(data)
        if pdf_bytes is not None:
            pdf_result = self.pdf_extractor.extract(pdf_bytes)
            result.tables.extend(pdf_result.tables)
            for t in result.tables:
                t.source = "docx-via-pdf"
            # Only keep vector-drawn figures from the PDF route; raster
            # images are already covered (at original quality) by the
            # direct zip extraction above.
            vector_figs = [f for f in pdf_result.figures if f.source == "vector-cluster"]
            for i, f in enumerate(vector_figs, start=len(result.figures)):
                f.index_on_page = i
            result.figures.extend(vector_figs)
            result.warnings.extend(pdf_result.warnings)
        else:
            result.warnings.append(
                "LibreOffice was unavailable; falling back to text-only table extraction "
                "(tables will not have pixel-accurate images)."
            )
            result.tables.extend(self._extract_tables_fallback(data))

        if not result.tables and not result.figures:
            result.warnings.append("No tables or figures were detected in this document.")

        return result

    # ------------------------------------------------------------------ #
    # Images straight from the .docx zip
    # ------------------------------------------------------------------ #
    def _extract_media_images(self, data: bytes) -> List[ExtractedFigure]:
        import io

        figures: List[ExtractedFigure] = []
        seen_hashes = set()
        idx = 0
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for info in zf.infolist():
                    if not info.filename.startswith("word/media/") or info.is_dir():
                        continue
                    raw = zf.read(info.filename)

                    ext = Path(info.filename).suffix.lower()
                    if ext in (".emf", ".wmf"):
                        # Legacy vector metafile formats Pillow/browsers can't
                        # render; skip rather than saving an unusable file.
                        continue

                    width, height = self._image_dimensions(raw)
                    # A decorative icon is usually both small in file size and
                    # small in pixels; a simple-but-real diagram can be small
                    # in one dimension without being small in the other (e.g.
                    # a clean vector-style PNG compresses very well but is
                    # still a full-size figure), so only drop images that are
                    # small on *both* axes.
                    is_small_bytes = len(raw) < self.min_image_bytes
                    is_small_pixels = width < 100 or height < 100
                    if is_small_bytes and is_small_pixels:
                        continue

                    h = sha256_bytes(raw)
                    if h in seen_hashes:
                        continue
                    seen_hashes.add(h)
                    figures.append(
                        ExtractedFigure(
                            page_number=0,
                            index_on_page=idx,
                            bbox=(0, 0, float(width or 0), float(height or 0)),
                            image_bytes=raw,
                            width=width or 0,
                            height=height or 0,
                            source="docx-media",
                        )
                    )
                    idx += 1
        except Exception as exc:  # pragma: no cover - defensive
            logger.warning("Failed reading DOCX media images: %s", exc)

        return figures

    def _image_dimensions(self, raw: bytes):
        try:
            from PIL import Image
            import io

            with Image.open(io.BytesIO(raw)) as img:
                return img.width, img.height
        except Exception:
            return 0, 0

    # ------------------------------------------------------------------ #
    # LibreOffice conversion
    # ------------------------------------------------------------------ #
    def _convert_to_pdf(self, data: bytes) -> Optional[bytes]:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            return None

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            src = tmp_path / "input.docx"
            src.write_bytes(data)
            try:
                subprocess.run(
                    [
                        soffice,
                        "--headless",
                        "--norestore",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(tmp_path),
                        str(src),
                    ],
                    check=True,
                    timeout=self.libreoffice_timeout,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
            except Exception as exc:
                logger.warning("LibreOffice conversion failed: %s", exc)
                return None

            out = tmp_path / "input.pdf"
            if not out.exists():
                return None
            return out.read_bytes()

    # ------------------------------------------------------------------ #
    # Fallback: text-only tables via python-docx (no LibreOffice)
    # ------------------------------------------------------------------ #
    def _extract_tables_fallback(self, data: bytes) -> List[ExtractedTable]:
        import io as _io

        from docx import Document

        from .utils import dedupe_column_names, normalize_rows, rows_look_like_table

        tables: List[ExtractedTable] = []
        try:
            doc = Document(_io.BytesIO(data))
        except Exception as exc:
            logger.warning("python-docx failed to open document: %s", exc)
            return tables

        for idx, table in enumerate(doc.tables):
            header, body, fills = _docx_table_header_and_body(table)
            if not header:
                continue
            rows = normalize_rows([header] + body)
            if not rows_look_like_table(rows):
                continue
            header, body = rows[0], rows[1:]

            df = pd.DataFrame(body, columns=dedupe_column_names(header))

            image_bytes = self._render_table_image(df, f"Table {idx + 1}", cell_fills=fills)
            tables.append(
                ExtractedTable(
                    page_number=0,
                    index_on_page=idx,
                    bbox=(0, 0, 0, 0),
                    dataframe=df,
                    image_bytes=image_bytes,
                    source="docx-text",
                )
            )

        return tables

    def _render_table_image(
        self,
        df: pd.DataFrame,
        title: str,
        cell_fills: Optional[List[List[Optional[str]]]] = None,
    ) -> bytes:
        """Draw the table with real text measurement so columns/rows autofit
        to their content and text never overlaps.

        matplotlib's ``ax.table`` cells have fixed heights and don't wrap
        text, which overlaps or clips anything longer than the cell -- so
        this draws directly with Pillow instead, using the TrueType fonts
        matplotlib already bundles (no new dependency) for exact pixel
        measurement of each line before laying out the grid.

        ``cell_fills`` optionally gives each body cell's own background
        color ("RRGGBB", row-major matching ``df``) -- used for tables
        (e.g. a Gantt-style schedule) whose real content is Word cell
        shading rather than text, so the rendered image stays faithful to
        the source instead of showing a blank grid.
        """
        import io as _io

        import matplotlib
        from PIL import Image, ImageDraw, ImageFont

        font_dir = Path(matplotlib.get_data_path()) / "fonts" / "ttf"
        font_size, title_font_size = 15, 20
        regular_font = ImageFont.truetype(str(font_dir / "DejaVuSans.ttf"), font_size)
        bold_font = ImageFont.truetype(str(font_dir / "DejaVuSans-Bold.ttf"), font_size)
        title_font = ImageFont.truetype(str(font_dir / "DejaVuSans-Bold.ttf"), title_font_size)

        columns = [str(c) for c in df.columns]
        rows = [["" if pd.isna(v) else str(v) for v in row] for row in df.itertuples(index=False, name=None)]
        n_cols = len(columns)

        pad_x, pad_y, line_gap = 12, 8, 4
        min_col_w, max_col_w = 90, 280

        measure_img = Image.new("RGB", (1, 1))
        measure_draw = ImageDraw.Draw(measure_img)

        def text_width(s: str, font: ImageFont.FreeTypeFont) -> float:
            return measure_draw.textlength(s, font=font)

        def line_height(font: ImageFont.FreeTypeFont) -> int:
            bbox = font.getbbox("Ag")
            return (bbox[3] - bbox[1]) + line_gap

        def wrap(text: str, font: ImageFont.FreeTypeFont, col_width: int) -> List[str]:
            avail = max(col_width - 2 * pad_x, 10)
            lines: List[str] = []
            for para in str(text).split("\n"):
                if not para:
                    lines.append("")
                    continue
                words = para.split(" ")
                current = words[0]
                for word in words[1:]:
                    candidate = f"{current} {word}"
                    if text_width(candidate, font) <= avail:
                        current = candidate
                    else:
                        lines.append(current)
                        current = word
                lines.append(current)
            return lines or [""]

        # Column widths: size to the longest line each column actually
        # contains (header or data), clamped to a sane min/max so one long
        # cell can't blow out the whole table -- it wraps instead.
        col_widths = []
        for c in range(n_cols):
            natural = max((text_width(part, bold_font) for part in columns[c].split("\n")), default=0)
            for row in rows:
                for part in row[c].split("\n"):
                    natural = max(natural, text_width(part, regular_font))
            col_widths.append(int(min(max(natural + 2 * pad_x, min_col_w), max_col_w)))

        header_lines = [wrap(columns[c], bold_font, col_widths[c]) for c in range(n_cols)]
        header_h = max((len(lines) for lines in header_lines), default=1) * line_height(bold_font) + 2 * pad_y

        body_lines: List[List[List[str]]] = []
        row_heights: List[int] = []
        reg_line_h = line_height(regular_font)
        for row in rows:
            cell_lines = [wrap(row[c], regular_font, col_widths[c]) for c in range(n_cols)]
            body_lines.append(cell_lines)
            row_heights.append(max((len(cl) for cl in cell_lines), default=1) * reg_line_h + 2 * pad_y)

        table_w = sum(col_widths)
        title_h = line_height(title_font) + 2 * pad_y if title else 0
        table_h = header_h + sum(row_heights)

        img = Image.new("RGB", (table_w + 1, title_h + table_h + 1), "white")
        draw = ImageDraw.Draw(img)

        y = 0
        if title:
            draw.text((table_w / 2, title_h / 2), title, font=title_font, fill="#1a1a1a", anchor="mm")
            y = title_h

        header_top = y
        draw.rectangle([0, y, table_w, y + header_h], fill=(46, 92, 138))
        x = 0
        for c in range(n_cols):
            cy = y + pad_y
            for line in header_lines[c]:
                draw.text((x + pad_x, cy), line, font=bold_font, fill="white")
                cy += line_height(bold_font)
            x += col_widths[c]
        y += header_h

        for r_idx, cell_lines in enumerate(body_lines):
            rh = row_heights[r_idx]
            x = 0
            for c in range(n_cols):
                fill = cell_fills[r_idx][c] if cell_fills else None
                cell_bg = f"#{fill}" if fill else ("#f2f2f2" if r_idx % 2 else "white")
                text_color = "white" if fill and not _is_light_hex(fill) else "#222222"
                draw.rectangle([x, y, x + col_widths[c], y + rh], fill=cell_bg)
                cy = y + pad_y
                for line in cell_lines[c]:
                    draw.text((x + pad_x, cy), line, font=regular_font, fill=text_color)
                    cy += reg_line_h
                x += col_widths[c]
            y += rh

        border = "#999999"
        x = 0
        for c in range(n_cols + 1):
            draw.line([(x, header_top), (x, header_top + table_h)], fill=border)
            if c < n_cols:
                x += col_widths[c]
        y = header_top
        draw.line([(0, y), (table_w, y)], fill=border)
        y += header_h
        draw.line([(0, y), (table_w, y)], fill=border)
        for rh in row_heights:
            y += rh
            draw.line([(0, y), (table_w, y)], fill=border)

        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf.read()
